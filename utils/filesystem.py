"""
File system utility functions for the gem-assist package.
These functions are used for file and directory operations.
"""

import os
import shutil
import glob
import platform
from typing import Dict, List, Any
import datetime
import mimetypes
import tempfile
from pydantic import BaseModel, Field

from .core import tool_message_print, tool_report_print
from gem import format_size

def get_current_directory() -> str:
    """
    Get the current working directory.

    Returns:
        str: The absolute path of the current working directory as a string.
    """
    tool_message_print("get_current_directory", [])
    try:
        return os.getcwd()
    except Exception as e:
        tool_report_print("Error getting current directory:", str(e), is_error=True)
        return f"Error getting current directory: {e}"

def list_dir(path: str, recursive: bool, files_only: bool, dirs_only: bool, bypass_safety: bool = False) -> list:
    """
    Returns a list of contents of a directory. It can handle listing files, directories, or both,
    and can do so recursively or not.

    Args:
        path: The path to the directory.
        recursive: Whether to list contents recursively. If True, it will traverse subdirectories.
        files_only: Whether to list only files. If True, directories are ignored.
        dirs_only: Whether to list only directories. If True, files are ignored.
        bypass_safety: If True, bypasses safety checks for dangerous paths. Use with caution.

    Returns:
        list: A list of dictionaries containing information about each item in the directory.
            Each dictionary has the keys:
            - 'name': The name of the file or directory.
            - 'path': The full path to the file or directory.
            - 'is_dir': A boolean indicating if the item is a directory.
            - 'size': The size of the file in a human-readable format (GB or MB), or 'N/A' for directories.
            
            Note that it can have different behavior based on given arguments, for example if you only need files, set `files_only=True` and ignore `dirs_only` and `recursive` arguments, they won't have any effect.
    """
    tool_message_print("list_dir", [("path", path), ("recursive", str(recursive)), 
                                   ("files_only", str(files_only)), ("dirs_only", str(dirs_only)), 
                                   ("bypass_safety", str(bypass_safety))])
    
    # Safety checks
    if recursive and not bypass_safety:
        # Check for dangerous paths that shouldn't be scanned recursively
        dangerous_paths = ["/", "/proc", "/sys", "/dev"]
        if path in dangerous_paths:
            tool_report_print("Warning:", f"Recursive listing of {path} is not allowed for safety reasons. Using non-recursive mode. (Use bypass_safety=True to override)", is_error=True)
            recursive = False
        
        # Apply item limit to prevent extremely long operations
        max_items = 1000  # Limit to 1000 items for recursive scans
    else:
        max_items = None  # No limit for non-recursive scans or when safety is bypassed
    
    items = []
    item_count = 0

    def add_item(item_path):
        nonlocal item_count
        if max_items is not None and item_count >= max_items:
            return False  # Stop collecting items if we hit the limit
        
        item_info = {
            'name': os.path.basename(item_path),
            'path': item_path,
            'is_dir': os.path.isdir(item_path),
            'size': format_size(os.path.getsize(item_path)) if os.path.isfile(item_path) else 'N/A'
        }
        items.append(item_info)
        item_count += 1
        return True  # Continue collecting items

    if recursive:
        try:
            for dirpath, dirnames, filenames in os.walk(path):
                if not files_only:
                    for dirname in dirnames:
                        if not add_item(os.path.join(dirpath, dirname)):
                            break
                
                if not dirs_only:
                    for filename in filenames:
                        if not add_item(os.path.join(dirpath, filename)):
                            break
                
                # Check if we've hit the limit after processing each directory
                if max_items is not None and item_count >= max_items:
                    tool_report_print("Warning:", f"Reached limit of {max_items} items. Output is truncated. (Use bypass_safety=True to see all items)", is_error=True)
                    break
        except PermissionError as e:
            tool_report_print("Error:", f"Permission denied: {e}", is_error=True)
    else:
        with os.scandir(path) as it:
            for entry in it:
                if files_only and entry.is_file():
                    add_item(entry.path)
                elif dirs_only and entry.is_dir():
                    add_item(entry.path)
                elif not files_only and not dirs_only:
                    add_item(entry.path)

    return items

def get_drives() -> list[dict]:
    """
    Get a list of drives on the system.

    Returns:
        list[dict]: A list of dictionaries containing information about each drive.
                     Each dictionary has the following keys:
                     - 'OsType': The OS type (e.g., "Windows", "Linux", "MacOS").
                     - 'Drive': The drive letter (e.g., "C:") or mount point (e.g., "/").
                     - 'Type': The drive type (e.g., "Fixed", "Removable", "Network").
                     - 'FileSystem': The file system type (e.e., "NTFS", "ext4", "apfs"), or 'N/A'.
                     - 'FreeSpace': The amount of free space in human-readable format (GB or MB), or 'N/A'.
                     - 'TotalSize': The total size of the drive in human-readable format (GB or MB), or 'N/A'.
    """
    tool_message_print("get_drives")
    drives = []
    os_type = platform.system()

    if os_type == "Windows":
        try:
            import wmi
            c = wmi.WMI()
            for drive in c.Win32_LogicalDisk():
                drive_type_map = {
                    0: "Unknown",
                    1: "No Root Directory",
                    2: "Removable",
                    3: "Fixed",
                    4: "Network",
                    5: "Compact Disc",
                    6: "RAM Disk"
                }
                drives.append({
                    'OsType': "Windows",
                    'Drive': drive.DeviceID,
                    'Type': drive_type_map.get(drive.DriveType, "Unknown"),
                    'FileSystem': drive.FileSystem if drive.FileSystem else 'N/A',
                    'FreeSpace': format_size(drive.FreeSpace) if drive.FreeSpace else 'N/A',
                    'TotalSize': format_size(drive.Size) if drive.Size else 'N/A'
                })
        except ImportError:
            tool_report_print("Warning:", "WMI module not available, Windows drive information will be limited", is_error=True)
            import psutil
            for partition in psutil.disk_partitions(all=False):
                try:
                    usage = psutil.disk_usage(partition.mountpoint)
                    drives.append({
                        'OsType': "Windows",
                        'Drive': partition.device,
                        'Type': "Unknown",
                        'FileSystem': partition.fstype if partition.fstype else 'N/A',
                        'FreeSpace': format_size(usage.free),
                        'TotalSize': format_size(usage.total)
                    })
                except (PermissionError, OSError):
                    # Skip partitions that can't be accessed
                    pass
    
    elif os_type == "Linux" or os_type == "Darwin":
        import psutil
        for partition in psutil.disk_partitions():
            try:
                disk_usage = psutil.disk_usage(partition.mountpoint)
                drives.append({
                    'OsType': os_type,
                    'Drive': partition.mountpoint,
                    'Type': partition.fstype,  # Filesystem type might serve as a decent "Type"
                    'FileSystem': partition.fstype if partition.fstype else 'N/A',
                    'FreeSpace': format_size(disk_usage.free),
                    'TotalSize': format_size(disk_usage.total)
                })
            except OSError:
                tool_report_print("Warning:", f"Failed to get drive information for {partition.mountpoint}. Skipping.", is_error=True)
    else:
        return []

    return drives

def get_directory_size(path: str, max_depth: int = 10, bypass_safety: bool = False) -> dict:
    """Get the size of the specified directory.

    Args:
      path: The path to the directory.
      max_depth: Maximum depth to traverse for size calculation (default: 10).
      bypass_safety: If True, bypasses safety checks and depth limits. Use with caution.

    Returns:
        dict: A dictionary containing the total size and the number of files in the directory.
        The dictionary has the following keys:
        - 'TotalSize': The total size of the directory in human-readable format (GB or MB).
        - 'FileCount': The number of files in the directory.
    """
    tool_message_print("get_directory_size", [("path", path), ("max_depth", str(max_depth)), 
                                             ("bypass_safety", str(bypass_safety))])
    total_size = 0
    file_count = 0
    
    # Safety check for dangerous paths
    dangerous_paths = ["/", "/proc", "/sys", "/dev"]
    if not bypass_safety and path in dangerous_paths:
        tool_report_print("Warning:", f"Getting size of {path} could be resource-intensive. Use bypass_safety=True to proceed.", is_error=True)
        return {
            'TotalSize': 'Size calculation skipped for safety',
            'FileCount': 'Count skipped for safety'
        }

    current_depth = 0
    for dirpath, dirnames, filenames in os.walk(path):
        # Check depth limit
        rel_path = os.path.relpath(dirpath, path)
        current_depth = 0 if rel_path == '.' else rel_path.count(os.sep) + 1
        
        if not bypass_safety and current_depth > max_depth:
            tool_report_print("Warning:", f"Max depth {max_depth} reached. Some files may not be counted.", is_error=True)
            continue
            
        for f in filenames:
            fp = os.path.join(dirpath, f)
            if os.path.isfile(fp):
                try:
                    total_size += os.path.getsize(fp)
                    file_count += 1
                except (OSError, PermissionError) as e:
                    tool_report_print("Error:", f"Could not access {fp}: {e}", is_error=True)

    return {
        'TotalSize': format_size(total_size),
        'FileCount': file_count
    }

def get_multiple_directory_size(paths: list[str], max_depth: int = 10, bypass_safety: bool = False) -> list[dict]:
    """Get the size of multiple directories.

    Args:
        paths: A list of paths to directories.
        max_depth: Maximum depth to traverse for size calculation (default: 10).
        bypass_safety: If True, bypasses safety checks and depth limits. Use with caution.

    Returns:
        list[dict]: A list of dictionaries containing the total size and the number of files in each directory.
        each item is the same as `get_directory_size`
    """
    tool_message_print("get_multiple_directory_size", [("paths", str(paths)), 
                                                      ("max_depth", str(max_depth)),
                                                      ("bypass_safety", str(bypass_safety))])
    return [get_directory_size(path, max_depth, bypass_safety) for path in paths]

def read_file(filepath: str) -> str:
    """
    Read content from a single file, in utf-8 encoding only. 
    !!!IMPORTANT: For non-text files (PDF, DOCX, XLSX, etc.) use read_file_content instead!!!
    This function is ONLY for plain text files like .txt, .py, etc.

    Args:
      filepath: The path to the file (must be a plain text file).

    Returns:
        str: The content of the file as a string.
    """
    tool_message_print("read_file", [("filepath", filepath)])
    try:
        with open(filepath, 'r', encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        tool_report_print("Error reading file:", str(e), is_error=True)
        return f"Error reading file: {e}"

def create_directory(paths: list[str]) -> bool:
    """
    Create single or multiple directories.

    Args:
      paths: A list of paths to the new directories.

    Returns:
        bool: True if directories were created successfully, False otherwise.
    """
    tool_message_print("create_directory", [("paths", str(paths))])
    try:
        success = True
        for path in paths:
            os.makedirs(path, exist_ok=True)
            tool_report_print("Created ✅:", path)
        return success
    except Exception as e:
        tool_report_print("Error creating directory:", str(e), is_error=True)
        return False

def get_file_metadata(filepath: str) -> dict:
    """
    Get metadata of a file.

    Args:
      filepath: The path to the file.

    Returns:
        dict: A dictionary containing file metadata:
              - 'creation_time': The timestamp of the file's creation.
              - 'modification_time': The timestamp of the file's last modification.
              - 'creation_time_readable': The creation time in ISO format.
              - 'modification_time_readable': The modification time in ISO format.
    """
    tool_message_print("get_file_metadata", [("filepath", filepath)])
    try:
        timestamp_creation = os.path.getctime(filepath)
        timestamp_modification = os.path.getmtime(filepath)
        return {
            'creation_time': timestamp_creation,
            'modification_time': timestamp_modification,
            'creation_time_readable': datetime.datetime.fromtimestamp(timestamp_creation).isoformat(),
            'modification_time_readable': datetime.datetime.fromtimestamp(timestamp_modification).isoformat()
        }
    except Exception as e:
        tool_report_print("Error getting file metadata:", str(e), is_error=True)
        return f"Error getting file metadata: {e}"

class FileData(BaseModel):
    file_path: str = Field(..., description="Path of the file, can be folder/folder2/filename.txt too")
    content: str = Field(..., description="Content of the file")

def write_files(files_data: list[FileData]) -> dict:
    """
    Write content to multiple files, supports nested directory file creation.
    
    Args:
      files_data: A list of FileData objects containing file paths and content.

    Returns:
      dict: A dictionary with file paths as keys and success status as values.
    """
    tool_message_print("write_files", [("count", str(len(files_data)))])
    results = {}
    
    for file_data in files_data:
        try:
            nested_dirs = os.path.dirname(file_data.file_path)
            if nested_dirs:
                os.makedirs(nested_dirs, exist_ok=True)

            with open(file_data.file_path, 'w', encoding="utf-8") as f:
                f.write(file_data.content)
            tool_report_print("Created ✅:", file_data.file_path)
            results[file_data.file_path] = True
        except Exception as e:
            tool_report_print("❌", file_data.file_path, is_error=True)
            tool_report_print("Error writing file:", str(e), is_error=True)
            results[file_data.file_path] = False

    success_count = sum(1 for success in results.values() if success)
    total_count = len(results)
    
    tool_report_print("Summary:", f"Wrote {success_count}/{total_count} files successfully")
    
    return results

def copy_file(src_filepath: str, dest_filepath: str) -> bool:
    """
    Copy a file from source to destination.

    Args:
      src_filepath: Path to the source file.
      dest_filepath: Path to the destination.

    Returns:
      bool: True if copy successful, False otherwise.
    """
    tool_message_print("copy_file", [("src_filepath", src_filepath), ("dest_filepath", dest_filepath)])
    try:
        shutil.copy2(src_filepath, dest_filepath) 
        tool_report_print("Status:", "File copied successfully")
        return True
    except Exception as e:
        tool_report_print("Error copying file:", str(e), is_error=True)
        return False

def move_file(src_filepath: str, dest_filepath: str) -> bool:
    """
    Move a file from source to destination.

    Args:
      src_filepath: Path to the source file.
      dest_filepath: Path to the destination.

    Returns:
      bool: True if move successful, False otherwise.
    """
    tool_message_print("move_file", [("src_filepath", src_filepath), ("dest_filepath", dest_filepath)])
    try:
        shutil.move(src_filepath, dest_filepath)
        tool_report_print("Status:", "File moved successfully")
        return True
    except Exception as e:
        tool_report_print("Error moving file:", str(e), is_error=True)
        return False
    
def rename_file(filepath: str, new_filename: str) -> bool:
    """
    Rename a file.

    Args:
      filepath: Current path to the file.
      new_filename: The new filename (not path, just the name).

    Returns:
      bool: True if rename successful, False otherwise.
    """
    tool_message_print("rename_file", [("filepath", filepath), ("new_filename", new_filename)])
    directory = os.path.dirname(filepath)
    new_filepath = os.path.join(directory, new_filename)
    try:
        os.rename(filepath, new_filepath)
        tool_report_print("Status:", "File renamed successfully")
        return True
    except Exception as e:
        tool_report_print("Error renaming file:", str(e), is_error=True)
        return False

def rename_directory(path: str, new_dirname: str) -> bool:
    """
    Rename a directory.

    Args:
      path: Current path to the directory.
      new_dirname: The new directory name (not path, just the name).

    Returns:
      bool: True if rename successful, False otherwise.
    """
    tool_message_print("rename_directory", [("path", path), ("new_dirname", new_dirname)])
    parent_dir = os.path.dirname(path)
    new_path = os.path.join(parent_dir, new_dirname)
    try:
        os.rename(path, new_path)
        tool_report_print("Status:", "Directory renamed successfully")
        return True
    except Exception as e:
        tool_report_print("Error renaming directory:", str(e), is_error=True)
        return False

def find_files(pattern: str, directory: str = ".", recursive: bool = False, 
               include_hidden: bool = False, max_results: int = 1000, 
               bypass_safety: bool = False) -> list[str]:
    """
    Searches for files (using glob) matching a given pattern within a specified directory.

    Args:
        pattern: The glob pattern to match (e.g., "*.txt", "data_*.csv").
        directory: The directory to search in (defaults to the current directory).
        recursive: Whether to search recursively in subdirectories (default is False).
        include_hidden: Whether to include hidden files (default is False).
        max_results: Maximum number of results to return (default is 1000).
        bypass_safety: If True, bypasses safety checks for dangerous paths and result limits. Use with caution.

    Returns:
        A list of file paths that match the pattern. Returns an empty list if no matches are found.
        Returns an appropriate error message if the directory does not exist or is not accessible.
    """
    tool_message_print("find_files", [("pattern", pattern), ("directory", directory), 
                                      ("recursive", str(recursive)), ("include_hidden", str(include_hidden)),
                                      ("max_results", str(max_results)), ("bypass_safety", str(bypass_safety))])
    try:
        if not os.path.isdir(directory):
            tool_report_print("Error:", f"Directory '{directory}' not found.", is_error=True)
            return f"Error: Directory '{directory}' not found."  # Clear error message
        
        # Safety checks for dangerous paths
        if recursive and not bypass_safety:
            dangerous_paths = ["/", "/proc", "/sys", "/dev"]
            if directory in dangerous_paths:
                tool_report_print("Warning:", f"Recursive searching in {directory} is not allowed for safety reasons. Use bypass_safety=True to override.", is_error=True)
                recursive = False
        
        full_pattern = os.path.join(directory, pattern)  # Combine directory and pattern
        matches = glob.glob(full_pattern, recursive=recursive, include_hidden=include_hidden)
        
        # Apply result limit unless bypassed
        if not bypass_safety and len(matches) > max_results:
            tool_report_print("Warning:", f"Found {len(matches)} matches, limiting to {max_results}. Use bypass_safety=True to see all results.", is_error=True)
            matches = matches[:max_results]

        # Check if the list is empty and return a message.
        if not matches:
            tool_report_print("Status:", "No files found matching the criteria.")
            return "No files found matching the criteria."

        tool_report_print("Status:", f"Found {len(matches)} matching files")
        return matches  # Return the list of matching file paths

    except OSError as e:
        tool_report_print("Error:", str(e), is_error=True)
        return f"Error: {e}"  # Return the system error message

def read_file_content(file_path: str, force_text_output: bool = False) -> Dict[str, Any]:
    """
    Read file content efficiently without leaving residual files.
    Automatically detects file type and uses the appropriate method.
    
    Args:
        file_path: Path to the file to read
        force_text_output: Whether to force output as text for all file types
        
    Returns:
        Dictionary containing file content and metadata
    """
    tool_message_print("read_file_content", [
        ("file_path", file_path),
        ("force_text_output", str(force_text_output))
    ])
    
    if not os.path.exists(file_path):
        tool_report_print("Error reading file:", f"File not found: {file_path}", is_error=True)
        return {"error": f"File not found: {file_path}"}
    
    file_ext = os.path.splitext(file_path)[1].lower()
    file_size = os.path.getsize(file_path)
    
    result = {
        "filename": os.path.basename(file_path),
        "file_size": file_size,
        "file_type": file_ext,
        "last_modified": datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")
    }
    
    try:
        # Text files - direct reading, always as text
        if file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv', '.log', '.sh']:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                result["content"] = content
                result["content_type"] = "text"
                result["encoding"] = "utf-8"
                
        # Excel files - use pandas without creating temp files
        elif file_ext in ['.xlsx', '.xls']:
            try:
                import pandas as pd
                # Import locally to prevent dependency issues
                from .document_utils import read_excel_file
                
                # Get Excel content as structured data
                if force_text_output:
                    # If text is requested, convert to CSV string
                    df = pd.read_excel(file_path)
                    result["content"] = df.to_csv(index=False)
                    result["content_type"] = "text"
                else:
                    # Return structured data
                    data = read_excel_file(file_path, return_format="dict")
                    result["content"] = data
                    result["content_type"] = "structured"
                    
                    # Add Excel-specific metadata
                    from .document_utils import read_excel_structure
                    result["structure"] = read_excel_structure(file_path)
            except ImportError:
                result["error"] = "pandas not installed, cannot read Excel files"
                result["content"] = f"Error: pandas not installed. Install with 'uv pip install pandas'"
                result["content_type"] = "error"
        
        # PDF files - read directly with PyPDF2, no temp files
        elif file_ext == '.pdf':
            try:
                # Import locally to prevent dependency issues
                from .document_utils import read_pdf_text
                content = read_pdf_text(file_path)
                
                result["content"] = content
                result["content_type"] = "text"
            except ImportError:
                result["error"] = "PyPDF2 not installed, cannot read PDF files"
                result["content"] = f"Error: PyPDF2 not installed. Install with 'uv pip install PyPDF2'"
                result["content_type"] = "error"
        
        # Word documents - read directly with python-docx
        elif file_ext in ['.docx', '.doc']:
            try:
                if file_ext == '.docx':
                    import docx
                    doc = docx.Document(file_path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    result["content"] = content
                    result["content_type"] = "text"
                else:
                    # For old .doc format, we need conversion, but will clean up after
                    with tempfile.TemporaryDirectory() as temp_dir:
                        from .document_utils import convert_document
                        converted = convert_document(file_path, "docx", temp_dir)
                        
                        if converted["success"]:
                            temp_file = converted["output_file"]
                            import docx
                            doc = docx.Document(temp_file)
                            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            result["content"] = content
                            result["content_type"] = "text"
                            # Temp files will be automatically deleted when the context manager exits
                        else:
                            result["error"] = f"Failed to convert .doc file: {converted.get('error')}"
                            result["content_type"] = "error"
            except ImportError:
                result["error"] = "python-docx not installed, cannot read Word documents"
                result["content"] = f"Error: python-docx not installed. Install with 'uv pip install python-docx'"
                result["content_type"] = "error"
        
        # Images - just report metadata, don't process unless requested
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']:
            result["content_type"] = "image"
            result["content"] = f"Image file: {os.path.basename(file_path)} ({file_size} bytes)"
            
            # Include image dimensions if PIL is available
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    result["width"] = img.width
                    result["height"] = img.height
                    result["mode"] = img.mode
            except ImportError:
                result["note"] = "Install Pillow for image dimensions: uv pip install Pillow"
        
        # Binary files - report metadata but don't try to read content
        else:
            # Try to guess content type
            mimetype, _ = mimetypes.guess_type(file_path)
            result["content_type"] = "binary"
            result["mime_type"] = mimetype or "application/octet-stream"
            result["content"] = f"Binary file: {os.path.basename(file_path)} ({file_size} bytes)"
            
            if force_text_output:
                # Try to read as text if explicitly requested
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read(4096)  # Just read a sample
                        result["content"] = content
                        result["note"] = "Binary file displayed as text (may contain unreadable characters)"
                except Exception as e:
                    result["error"] = f"Cannot display binary file as text: {str(e)}"
        
        tool_report_print("File read successfully:", 
                         f"Read {file_size} bytes from {os.path.basename(file_path)}")
        return result
        
    except Exception as e:
        tool_report_print("Error reading file:", str(e), is_error=True)
        result["error"] = str(e)
        result["content_type"] = "error"
        return result
